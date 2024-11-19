# -*- coding: utf-8 -*-
# Created on Wed Oct 21 17:18:07 2020

# @author: Davide Laghi

# Copyright 2021, the JADE Development Team. All rights reserved.

# This file is part of JADE.

# JADE is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# JADE is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with JADE.  If not, see <http://www.gnu.org/licenses/>.

#################### Modification by Matteo 18/11/2024##################################
########################################################################################
# The current script is a modification of the original script by Davide Laghi 
# It include the following changes:

# All the experimental benchmark have been deleted but the Okatvian one.

# Following the output.py script, the ExpAbstractOutput class has been created
# even if it still require major modifications.

# For the moment the classes such as SpectrumOutput, ShieldingOutput and
# MultispectrumOutput have been keeped so that the codes runs without errors, but 
# they must be editated appropriately in order to include other experimental benchmarks. 
########################################################################################

from __future__ import annotations

import abc
import json
import math
import os
import string
import re
import shutil
from abc import abstractmethod

import numpy as np
import pandas as pd
from docx.shared import Inches
from f4enix.input.MCNPinput import D1S_Input
from f4enix.output.mctal import Tally
from scipy.interpolate import interp1d
from tqdm import tqdm

import jade.atlas as at
from jade.output import MCNPBenchmarkOutput, MCNPSimOutput
from jade.plotter import Plotter
from jade.status import EXP_TAG

# RED color
CRED = "\033[91m"
CEND = "\033[0m"

FOILS_REACTION = {
    "Al": "Al(n,alpha)",
    "S": "S-32(n,p)",
    "In": "In-115(n,n')",
    "Rh": "Rh-103(n,n')",
    "Au": "Au-197(n,gamma)",
}

MCNP_UNITS = {"Energy": "MeV", "Time": "shakes"}

TALLY_NORMALIZATION = {
    "Oktavian": "lethargy",
    "ASPIS-PCA-Replica_flux": "lethargy",
}

ACTIVATION_REACTION = {
    "Ni-n2n": "Ni-58(n,2n)Ni-57",
    "Al": "Al-27(n,a)Na-24",
    "Fe": "Fe-56(n,p)Mn-56",
    "Ni-np": "Ni-58(n,p)Co-58",
    "Nb": "Nb-93(n,2n)Nb-92",
    "In": "In-115(n,n')In-115m",
    "Mn": "Mn-55(n,g)Mn-56",
    "Au": "Au-197(n,g)Au-198",
    "Rh": "Rh-103(n,n')Rh-103*",
    "S": "S-32(n,p)P-32",
    "Zr": "Zr-90(n,2n)Zr-89",
}

class ExpAbstractOutput(abc.ABC):
    @abc.abstractmethod
    def single_postprocess(self):
        """
        To be executed when a single pp is requested
        """
        pass

    @abc.abstractmethod
    def compare(self):
        """
        To be executed when a comparison is requested
        """

    @staticmethod
    def _get_output_files(results_path):
        """
        Recover the meshtal and outp file from a directory

        Parameters
        ----------
        results_path : str or path
            path where the MCNP results are contained.

        Raises
        ------
        FileNotFoundError
            if either meshtal or outp are not found.

        Returns
        -------
        mfile : path
            path to the meshtal file
        ofile : path
            path to the outp file

        """
        # Get mfile
        mfile = None
        ofile = None

        for file in os.listdir(results_path):
            if file[-1] == "m":
                mfile = file
            elif file[-1] == "o":
                ofile = file

        if mfile is None or ofile is None:
            raise FileNotFoundError(
                """
 The followig path does not contain either the .m or .o file:
 {}""".format(
                    results_path
                )
            )

        mfile = os.path.join(results_path, mfile)
        ofile = os.path.join(results_path, ofile)

        return mfile, ofile


class ExpBenchmarkOutput(ExpAbstractOutput):
    def __init__(self, lib: str, code: str, testname: str, session: Session, *args, **kwargs):
        """
        This extends the Benchmark Output and creates an abstract class
        for all experimental outputs.
        Parameters
        ----------
        *args : TYPE
            see BenchmarkOutput doc.
        **kwargs : TYPE
            see BenchmarkOutput doc.
        multiplerun : bool
            this additional keyword specifies if the benchmark is composed
            by more than one MCNP run. It defaults to False.
        Returns
        -------
        None.

        """
        self.raw_data = {}  # Raw data
        self.outputs = {}  # outputs linked to the benchmark
        self.testname = testname  # test name
        self.code_path = os.getcwd()  # path to code
        self.state = session.state
        self.session = session
        self.path_templates = session.path_templates
        # Add a special keyword for experimental benchmarks
        try:
            multiplerun = kwargs.pop("multiplerun")
        except KeyError:
            # Default to False
            multiplerun = False
        # Recover session and testname
        session = args[3]
        testname = args[2]
        super().__init__(*args, **kwargs)
        # The experimental data needs to be loaded
        self.path_exp_res = os.path.join(session.path_exp_res, testname)

        # Updated to handle multiple codes
        # initialize them so that intellisense knows they are available
        self.mcnp = False
        self.d1s = False
        self.serpent = False
        self.d1s = False
        for available_code in CODES.values():
            if code == available_code:
                setattr(self, available_code, True)
                self.raw_data[code] = {}
                self.outputs[code] = {}
            else:
                setattr(self, available_code, False)

        self.code = code  # this can be handy in a lot of places to avoid if else

        # Add the raw path data (not created because it is a comparison)
        out = os.path.dirname(self.atlas_path)
        raw_path = os.path.join(out, "Raw_Data")
        if not os.path.exists(raw_path):
            os.mkdir(raw_path)
        self.raw_path = raw_path
        self.multiplerun = multiplerun

        # Read the metadata from the simulations
        metadata = {}
        for lib, test_path in self.test_path.items():
            if lib == EXP_TAG:
                continue
            code = args[1]
            if self.multiplerun:
                # I still need only one metadata. They should be all the same
                results_path = os.path.join(test_path, os.listdir(test_path)[0], code)
                metadata_lib = self._read_metadata_run(results_path)
            else:
                results_path = os.path.join(test_path, code)
                self.metadata_lib = self._read_metadata_run(results_path)
            metadata[lib] = metadata_lib
        self.metadata = metadata

    def single_postprocess(self) -> None:
        """
        Always raise an Attribute Error since no single post-processing is
        foreseen for experimental benchmarks
        Raises
        ------
        AttributeError
            DESCRIPTION.
        Returns
        -------
        None.
        """
        raise AttributeError("\n No single pp is foreseen for exp benchmark")

    def compare(self) -> None:
        """
        Complete the routines that perform the comparison of one or more
        libraries results with the experimental ones.
        Returns
        -------
        None.
        """
        print(" Exctracting outputs...")
        self._extract_outputs()

        print(" Read experimental results....")
        self._read_exp_results()

        print(" Dumping raw data...")
        self._print_raw()

        print(" Generating Excel Recap...")
        self.pp_excel_comparison()

        print(" Creating Atlas...")
        self.build_atlas()

    def pp_excel_comparison(self) -> None:
        """
        At the moment everything is handled by _pp_excel_comparison that needs
        to be implemented in each child class. Some standard procedures may be
        added in the feature in order to reduce the amount of ex-novo coding
        necessary to implement a new experimental benchmark.
        Returns
        -------
        None.
        """
        self._pp_excel_comparison()

    def build_atlas(self) -> None:
        """
        Creation and saving of the atlas are handled by this function while
        the actual filling of the atlas is left to _build_atlas which needs
        to be implemented for each child class.
        Returns
        -------
        None.
        """
        # Build a temporary folder for images
        tmp_path = os.path.join(self.atlas_path, "tmp")

        os.mkdir(tmp_path)

        globalname = ""
        for lib in self.lib:
            globalname = globalname + lib + "_Vs_"
        globalname = globalname[:-4]
        globalname = self.testname + "_" + globalname
        # Initialize the atlas
        template = os.path.join(self.session.path_templates, "AtlasTemplate.docx")
        atlas = at.Atlas(template, globalname)

        # Fill the atlas
        atlas = self._build_atlas(tmp_path, atlas)

        atlas.save(self.atlas_path)
        # Remove tmp images
        shutil.rmtree(tmp_path)

    def _extract_single_output(
        self, results_path: str | os.PathLike, folder: str, lib: str
    ) -> tuple[pd.DataFrame, str]:
        """Method to extract single output data from MCNP files

        Parameters
        ----------
        results_path : str | os.PathLike
            Path to simulations results.
        folder : str
            Sub-folder for multiple run case.
        lib : str
            Test library.

        Returns
        -------
        tallydata : pd.DataFrame
            Pandas dataframe containing tally data.
        input : str
            Test name.
        """
        mfile, ofile, meshtalfile = self._get_output_files(results_path)
        # Parse output
        output = MCNPSimOutput(mfile, ofile, meshtalfile)

        # need to extract the input in case of multi
        if self.multiplerun:
            pieces = folder.split("_")
            input = pieces[-1]
            if input not in self.inputs:
                self.inputs.append(input)
            self.outputs[input, lib] = output
            # Get the meaningful results
            self.results[input, lib] = self._processMCNPdata(output)
        else:
            # just treat it as a special case of multiple run
            self.outputs[self.testname, lib] = output
            # Get the meaningful results
            self.results[self.testname, lib] = self._processMCNPdata(output)
            input = self.testname

        return output.tallydata, input

    def _extract_outputs(self) -> None:
        """
        Extract, organize and store the results coming from the different codes
        runs

        Returns
        -------
        None.
        """
        self.outputs = {}
        self.results = {}

        # Each output object is processing only one code at the time at the moment
        if self.mcnp:
            code_tag = "mcnp"
        if self.openmc:
            print("Experimental comparison not implemented for OpenMC")
            return
        if self.serpent:
            print("Experimental comparison not implemented for Serpent")
            return
        if self.d1s:
            code_tag = "d1s"

        # only multiple runs have multiple inputs
        if self.multiplerun:
            self.inputs = []
        else:
            self.inputs = [self.testname]

        # Iterate on the different libraries results except 'Exp'
        for lib, test_path in self.test_path.items():
            if lib != EXP_TAG:
                if self.multiplerun:
                    # Results are organized by folder and lib
                    code_raw_data = {}
                    for folder in os.listdir(test_path):
                        results_path = os.path.join(test_path, folder, code_tag)
                        tallydata, input = self._extract_single_output(
                            results_path, folder, lib
                        )
                        code_raw_data[input, lib] = tallydata

                # Results are organized just by lib
                else:
                    results_path = os.path.join(test_path, code_tag)
                    tallydata, input = self._extract_single_output(
                        results_path, None, lib
                    )
                    code_raw_data = {(self.testname, lib): tallydata}

                # Adjourn raw Data
                # self.raw_data[code_tag].update(code_raw_data)
                self.raw_data.update(code_raw_data)

    def _read_exp_results(self) -> None:
        """
        Read all experimental results and organize it in the self.exp_results
        dictionary.
        If multirun is set to true the first layer of the dictionary will
        consist in the different folders and the second layer will be the
        different files. If it is not multirun, insetead, only one layer of the
        different files will be generated.
        All files need to be in .csv format. If a more complex format is
        provided, the user should ovveride the _read_exp_file method.
        Returns
        -------
        None.
        """
        exp_results = {}
        if self.multiplerun:
            # Iterate on each folder and then in each file, read them and
            # build the result dic
            for folder in os.listdir(self.path_exp_res):
                exp_results[folder] = {}
                cp = os.path.join(self.path_exp_res, folder)
                for file in os.listdir(cp):
                    filename = file.split(".")[0]
                    filepath = os.path.join(cp, file)
                    df = self._read_exp_file(filepath)
                    c = df.columns.tolist()[1]
                    df = df[df[c] > 2e-38]
                    exp_results[folder][filename] = df
        else:
            # Iterate on each each file, read it and
            # build the result dic
            exp_results[self.testname] = {}
            for file in os.listdir(self.path_exp_res):
                filename = file.split(".")[0]
                filepath = os.path.join(self.path_exp_res, file)
                df = self._read_exp_file(filepath)
                c = df.columns.tolist()[1]
                df = df[df[c] > 2e-38]
                exp_results[self.testname][filename] = df

        self.exp_results = exp_results

    @staticmethod
    def _read_exp_file(filepath: str | os.PathLike) -> pd.DataFrame:
        """
        Default way of reading a csv file
        Parameters
        ----------
        filepath : path/str
            experimental file results to be read.
        Returns
        -------
        pd.DataFrame
            Contain the data read.
        """
        return pd.read_csv(filepath)

    def _print_raw(self) -> None:
        """
        Dump all the raw data
        Returns
        -------
        None.
        """
        raw_to_print = self.raw_data.items()

        for (folder, lib), item in raw_to_print:
            # Create the lib directory if it is not there
            cd_lib = os.path.join(self.raw_path, lib)
            if not os.path.exists(cd_lib):
                os.mkdir(cd_lib)
                # dump also the metadata if it is the first time
                with open(
                    os.path.join(cd_lib, "metadata.json"), "w", encoding="utf-8"
                ) as f:
                    json.dump(self.metadata[lib], f)

            # Dump everything
            for key, data in item.items():
                if folder == self.testname:
                    file = os.path.join(cd_lib, str(key) + ".csv")
                else:
                    file = os.path.join(cd_lib, folder + " " + str(key) + ".csv")
                data.to_csv(file, header=True, index=False)

    @abstractmethod
    def _processMCNPdata(self, output: MCNPSimOutput):
        """
        Given an mctal file object return the meaningful data extracted. Some
        post-processing on the data may be foreseen at this stage.
        Parameters
        ----------
        output : MCNPoutput
            object representing an MCNP output.
        Returns
        -------
        item :
            the type of item can vary based on what the user intends to do
            whith it. It will be stored in an organized way in the self.results
            dictionary
        """
        item = None
        return item

    @abstractmethod
    def _pp_excel_comparison(self) -> None:
        """
        Responsible for producing excel outputs
        Returns
        -------
        """
        pass

    @abstractmethod
    def _build_atlas(self, tmp_path: str | os.PathLike, atlas: at.Atlas) -> at.Atlas:
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.
        Parameters
        ----------
        tmp_path : path
            path to the temporary folder where to dump images.
        atlas : Atlas
            Object representing the plot Atlas.
        Returns
        -------
        atlas : Atlas
            After being filled the atlas is returned.
        """
        atlas = None
        return atlas




class SpectrumOutput(ExpBenchmarkOutput):

    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str | os.PathLike
            path to the temporary folder containing the plots for the atlas
        atlas : at.Atlas
            Object representing the plot Atlas.

        Returns
        -------
        atlas : at.Atlas
            Object representing the plot Atlas.
        """
        self.tables = []
        self.bench_conf = pd.read_excel(self.cnf_path)
        self.bench_conf = self.bench_conf.set_index(["Tally"])
        # Loop over benchmark cases
        for input in tqdm(self.inputs, desc=" Inputs: "):
            # Loop over tallies
            for tally in self.outputs[(input, self.lib[1])].mctal.tallies:
                # Get tally number and info
                tallynum, particle, xlabel = self._get_tally_info(tally)
                # Collect data
                quantity_CE = self.bench_conf.loc[tallynum, "Y Label"]
                e_int = self.bench_conf.loc[tallynum, "C/E X Quantity intervals"]
                e_int = e_int.split("-")

                # Convert the list of number strings into a list of integers
                e_intervals = [float(num) for num in e_int]
                data, xlabel = self._data_collect(
                    input, str(tallynum), quantity_CE, e_intervals
                )
                if not data:
                    continue

                # Use re.findall to extract all substrings between '[' and ']'
                unit = self.bench_conf.loc[tallynum, "Y Unit"]
                quantity = self.bench_conf.loc[tallynum, "Quantity"]
                title = self._define_title(input, quantity_CE)
                atlas.doc.add_heading(title, level=1)
                # Once the data is collected it is passed to the plotter
                outname = "tmp"
                plot = Plotter(
                    data,
                    title,
                    tmp_path,
                    outname,
                    quantity,
                    unit,
                    xlabel,
                    self.testname,
                )
                img_path = plot.plot("Experimental points")
                # Insert the image in the atlas
                atlas.insert_img(img_path)

        # Dump C/E table
        self._dump_ce_table()

        return atlas

    def _get_tally_info(self, tally: Tally) -> tuple[int, str, str]:
        """
        Extracts and assigns information from the tally object, as well as
        information from the benchmark config variable

        Parameters
        ----------
        tally : Tally
            F4Enix tally object

        Returns
        -------
        tallynum : int
            Tally number of the tally being plotted
        particle : str
            Type of quantity being plotted on the X axis
        quant + unit : str
            Unit of quantity being plotted on the X axis
        """
        tallynum = tally.tallyNumber
        particle = tally.particleList[np.where(tally.tallyParticles == 1)[0][0]]
        quant = self.bench_conf.loc[tallynum, "X Quantity"]
        unit = self.bench_conf.loc[tallynum, "X Unit"]
        return tallynum, particle, quant + " [" + unit + "]"

    def _define_title(self, input: str, quantity_CE: str) -> str:
        """Assigns the title for atlas plot

        Parameters
        ----------
        input : str
            String containing the name of the benchmark case being run, will
            be subfolder test name for benchmarks with multiple runs
        quantity_CE : str
            String containing the Y axis variable from benchmark config file

        Returns
        -------
        title : str
            Title string
        """

        if not self.multiplerun:
            title = self.testname + ", " + quantity_CE
        else:
            title = self.testname + " " + input + ", " + quantity_CE
        return title

    def _dump_ce_table(self) -> None:
        """
        Generates the C/E table and dumps them as an .xlsx file

        Returns
        -------
        None
        """
        final_table = pd.concat(self.tables)
        skipcol_global = 0
        binning_list = ["Energy", "Time"]
        for x_ax in binning_list:  # to update if other binning will be used
            x_lab = x_ax[0]
            col_check = "Max " + x_lab
            ft = final_table.set_index(["Input"])

            if col_check not in final_table.columns.tolist():
                continue
            else:
                todump = final_table.set_index(["Input", "Quantity", "Library"])
            for binning in binning_list:
                if binning == x_ax:
                    continue
                else:
                    # if tallies only have one type of binning KeyError could
                    # arise
                    try:
                        todump = todump.drop(
                            columns=["Min " + binning[0], "Max " + binning[0]]
                        )
                        ft = ft.drop(columns=["Min " + binning[0], "Max " + binning[0]])
                    except KeyError:
                        continue

            todump = todump.dropna(subset=["Max " + x_lab])
            ft = ft.dropna(subset=["Max " + x_lab])
            ex_outpath = os.path.join(
                self.excel_path, self.testname + "_" + x_ax + "_CE_tables.xlsx"
            )

            # Create a Pandas Excel writer using XlsxWriter as the engine.
            with pd.ExcelWriter(ex_outpath, engine="xlsxwriter") as writer:
                # dump global table
                todump = todump[
                    [
                        "Min " + x_lab,
                        "Max " + x_lab,
                        "C/E",
                        "Standard Deviation (σ)",
                    ]
                ]

                todump.to_excel(writer, sheet_name="Global")
                col_min = x_lab + "-min " + "[" + MCNP_UNITS[x_ax] + "]"
                col_max = x_lab + "-max " + "[" + MCNP_UNITS[x_ax] + "]"
                # Elaborate table for better output format

                ft[col_min] = ft["Min " + x_lab]
                ft[col_max] = ft["Max " + x_lab]

                ft["C/E (mean +/- σ)"] = (
                    ft["C/E"].round(2).astype(str)
                    + " +/- "
                    + ft["Standard Deviation (σ)"].round(2).astype(str)
                )
                # Delete all confusing columns
                for column in [
                    "Min " + x_lab,
                    "Max " + x_lab,
                    "C/E",
                    "Standard Deviation (σ)",
                ]:
                    del ft[column]

                # Dump also table material by material
                for input in self.inputs:
                    # dump material table
                    todump = ft.loc[input]

                    todump = todump.pivot(
                        index=["Quantity", col_min, col_max],
                        columns="Library",
                        values="C/E (mean +/- σ)",
                    )

                    todump.sort_values(by=[col_min])

                    todump.to_excel(writer, sheet_name=input, startrow=2)
                    ws = writer.sheets[input]
                    if skipcol_global == 0:
                        ws.write_string(0, 0, '"C/E (mean +/- σ)"')

                    # adjust columns' width
                    writer.sheets[input].set_column(0, 4, 18)

        return

    def _data_collect(
        self, input: str, tallynum: str, quantity_CE: str, e_intervals: list
    ) -> tuple[list, str]:
        """Collect data for C/E tables

        Parameters
        ----------
        input : str
            String containing the name of the benchmark case being run, will
            be subfolder test name for benchmarks with multiple runs
        tallynum : int
            tally number to be printed to table
        quantity_CE : str
            String containing the Y axis variable from benchmark config file
        e_intervals : list
            list of energy intervals from experimental benchmark config file

        Returns
        -------
        data : list
            list of dictionaries containing the data to be printed to the table
        x_lab : str
            Name of quantity being compared (not utilised?)

        """
        if self.multiplerun:
            filename = self.testname + "_" + input + "_" + str(tallynum)
        else:
            filename = self.testname + "_" + str(tallynum)
        # check if correspondent experimental data exists
        try:
            col_idx = self.exp_results[input][filename].columns.tolist()
        except KeyError:
            return None, None

        x_lab = col_idx[0]
        y_lab = col_idx[1]
        x = self.exp_results[input][filename][col_idx[0]].values
        y = self.exp_results[input][filename][col_idx[1]].values
        err = self.exp_results[input][filename][col_idx[2]].values
        # lib will be passed to the plotter
        lib = {"x": x, "y": y, "err": err, "ylabel": "Experiment"}
        # Get also the interpolator
        interpolator = interp1d(x, y, fill_value=0, bounds_error=False)
        # Collect the data to be plotted
        data = [lib]  # The first one should be the exp one
        for lib_tag in self.lib[1:]:  # Avoid exp
            lib_name = self.session.conf.get_lib_name(lib_tag)
            try:  # The tally may not be defined
                # Data for the plotter
                values = self.results[input, lib_tag][tallynum]
                lib = {
                    "x": values[x_lab],
                    "y": values["C"],
                    "err": values["Error"],
                    "ylabel": lib_name,
                }
                data.append(lib)
                # data for the table
                table = _get_tablevalues(
                    values, interpolator, x=x_lab, e_intervals=e_intervals
                )
                table["Quantity"] = quantity_CE
                table["Input"] = input
                table["Library"] = lib_name
                self.tables.append(table)
            except KeyError:
                # The tally is not defined
                pass
        return data, x_lab

    def _pp_excel_comparison(self) -> None:
        """
        Excel is actually printed by the build atlas in this case

        Returns
        -------
        None
        """
        # Excel is actually printed by the build atlas in this case
        pass

    def _processMCNPdata(self, output: MCNPSimOutput) -> dict:
        """
        given the mctal file the lethargy flux and energies are returned
        both for the neutron and photon tally

        Parameters
        ----------
        output : MCNPSimOutput
            object representing the MCNP output.
        Returns
        -------
        res : dict
            contains the extracted lethargy flux and energies.
        """

        res = {}
        # Read tally energy binned fluxes
        for tallynum, data in output.tallydata.items():
            tallynum = str(tallynum)
            res2 = res[tallynum] = {}
            x_axis = data.columns.tolist()[0]

            # Delete the total value
            data = data.set_index(x_axis).drop("total").reset_index()
            flux, energies, errors = self._parse_data_df(data, output, x_axis, tallynum)

            res2[x_axis + " [" + MCNP_UNITS[x_axis] + "]"] = energies
            res2["C"] = flux
            res2["Error"] = errors

            res[tallynum] = res2

        return res

    def _parse_data_df(
        self, data: pd.DataFrame, output: MCNPSimOutput, x_axis: str, tallynum: str
    ) -> tuple[list, list, list]:
        """
        Read information from data DataFrame

        Parameters
        ----------
        data : pd.DataFrame
            DataFrame containing all tally data for an output
        output : MCNPSimOutput
            MCNP output object generated by MCNP parser
        x_axis : str
            X axis title
        tallynum : int
            Tally number, used to determine behaviour for protons and
            neutrons

        Returns
        -------
        flux : list
            list of binned flux values
        energies: list
            list of energy bin boundaries
        errors : list
            list of binned error values
        """
        # Generate a folder for each library
        flux = data["Value"].values
        energies = data[x_axis].values
        errors = data["Error"].values

        if TALLY_NORMALIZATION[self.testname] == "lethargy":
            # Energies for lethargy computation
            ergs = [1e-10]  # Additional "zero" energy for lethargy computation
            ergs.extend(energies.tolist())
            ergs = np.array(ergs)

            # Different behaviour for photons and neutrons
            for tally in output.mctal.tallies:
                if tallynum == str(tally.tallyNumber):
                    particle = tally.particleList[
                        np.where(tally.tallyParticles == 1)[0][0]
                    ]
            if particle == "Neutron":
                flux = flux / np.log((ergs[1:] / ergs[:-1]))
            elif particle == "Photon":
                flux = flux / (ergs[1:] - ergs[:-1])

        elif TALLY_NORMALIZATION[self.testname] == "energy bins":
            # Energies for lethargy computation
            data["bin"] = None

            prev_e = 0

            for e in data[x_axis].unique().tolist():
                data.loc[data[x_axis] == e, "bin"] = e - prev_e
                prev_e = e
            flux = flux / data["bin"].values
        return flux, energies, errors


def _get_tablevalues(
    df: pd.DataFrame,
    interpolator: function,
    x: str = "Energy [MeV]",
    y: str = "C",
    e_intervals: list = [0.1, 1, 5, 10, 20],
):
    """
    Given the benchmark and experimental results returns a df to compile the
    C/E table for energy intervals

    Parameters
    ----------
    df : dict
        benchmark data.
    interpolator : func
        interpolator from experimental data.
    x : str, optional
        x column. The default is 'Energy [MeV]'.
    y : str, optional
        y columns. The default is 'C'.
    e_intervals : list, optional
        energy intervals to be used. The default is [0, 0.1, 1, 5, 10, 20].

    Returns
    -------
    pd.DataFrame
        C/E table per energy interval.

    """
    rows = []
    df = pd.DataFrame(df)
    df["Exp"] = interpolator(df[x])
    df["C/E"] = df[y] / df["Exp"]
    # it is better here to drop inf values because it means that the
    # interpolated experiment was zero, i.e., no value available
    df.replace([np.inf, -np.inf], np.nan, inplace=True)  # replace inf with NaN
    df.dropna(subset=["C/E"], how="all", inplace=True)  # drop the inf rows

    e_min = e_intervals[0]
    for e_max in e_intervals[1:]:
        red = df[e_min < df[x]]
        red = red[red[x] < e_max]
        mean = red["C/E"].mean()
        std = red["C/E"].std()
        row = {
            "C/E": mean,
            "Standard Deviation (σ)": std,
            "Max " + x[0]: e_max,
            "Min " + x[0]: e_min,
        }
        rows.append(row)
        # adjourn min energy
        e_min = e_max

    return pd.DataFrame(rows)





class ShieldingOutput(ExpBenchmarkOutput):

        Parameters
        ----------
        output : MCNPSimOutput
            MCNP simulation output object

        Returns
        -------
        None
        """
        return None

    def _pp_excel_comparison(self) -> None:
        """
        This method prints C/E tables for shielding benchmark comparisons

        Returns
        -------
        None.
        """
        # FNG SiC specific corrections/normalisations
        fngsic_k = [0.212, 0.204, 0.202, 0.202]  # Neutron sensitivity of TL detectors
        fngsic_norm = 1.602e-13 * 1000  # J/MeV * g/kg
        lib_names_dict = {}
        column_names = []
        column_names.append(("Exp", "Value"))
        column_names.append(("Exp", "Error"))
        for lib in self.lib[1:]:
            namelib = self.session.conf.get_lib_name(lib)
            lib_names_dict[namelib] = lib
            column_names.append((namelib, "Value"))
            column_names.append((namelib, "C/E"))
            column_names.append((namelib, "C/E Error"))

        names = ["Library", ""]
        column_index = pd.MultiIndex.from_tuples(column_names, names=names)
        # filepath = self.excel_path_mcnp + '\\' + self.testname + '_CE_tables.xlsx'
        filepath = os.path.join(self.excel_path, f"{self.testname}_CE_tables.xlsx")
        with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
            # TODO Replace when other transport codes implemented.
            code = "mcnp"
            for mat in self.inputs:
                exp_folder = os.path.join(self.path_exp_res, mat)
                exp_filename = self.testname + "_" + mat + ".csv"
                exp_filepath = os.path.join(exp_folder, exp_filename)
                exp_data_df = pd.read_csv(exp_filepath)
                # Get experimental data and errors for the selected benchmark case
                x = exp_data_df["Depth"].values.tolist()
                indexes = pd.Index(data=x, name="Depth [cm]")
                df_tab = pd.DataFrame(index=indexes, columns=column_index)
                for idx_col in df_tab.columns.values.tolist():
                    if idx_col[0] == "Exp":
                        if idx_col[1] == "Value":
                            vals = exp_data_df.loc[:, "Reaction Rate"].tolist()
                            df_tab[idx_col] = vals
                        else:
                            vals = exp_data_df.loc[:, "Error"].to_numpy() / 100
                            vals = vals.tolist()
                            df_tab[idx_col] = vals
                    else:
                        t = (mat, lib_names_dict[idx_col[0]])
                        if idx_col[1] == "Value":
                            if mat != "TLD":
                                vals = self.raw_data[t][4]["Value"].values[: len(x)]
                            else:
                                # FNG SiC experiment measured the total dose
                                if self.testname == "FNG-SiC":
                                    # Neutron dose
                                    Dn = (
                                        self.raw_data[t][16]["Value"].values[: len(x)]
                                    ) * fngsic_norm
                                    Dn_multiplied = [
                                        value * constant
                                        for value, constant in zip(Dn, fngsic_k)
                                    ]
                                    # Photon dose
                                    Dp = (
                                        self.raw_data[t][26]["Value"].values[: len(x)]
                                    ) * fngsic_norm
                                    # Sum neutron and photon dose with neutron sensitivity as a function of depth
                                    Dt = [sum(pair) for pair in zip(Dn_multiplied, Dp)]
                                    vals = Dt
                                else:
                                    vals = self.raw_data[t][6]["Value"].values[: len(x)]
                            df_tab[idx_col] = vals
                        elif idx_col[1] == "C/E Error":
                            if mat != "TLD":
                                errs = self.raw_data[t][4]["Error"].values[: len(x)]
                            else:
                                if self.testname == "FNG-SiC":
                                    errs = np.sqrt(
                                        np.square(
                                            self.raw_data[t][16]["Error"].values[
                                                : len(x)
                                            ]
                                        )
                                        + np.square(
                                            self.raw_data[t][26]["Error"].values[
                                                : len(x)
                                            ]
                                        )
                                    )
                                else:
                                    errs = self.raw_data[t][6]["Error"].values[: len(x)]
                            vals1 = np.square(errs)
                            vals2 = np.square(
                                exp_data_df.loc[:, "Error"].to_numpy() / 100
                            )
                            ce_err = np.sqrt(vals1 + vals2)
                            ce_err = ce_err.tolist()
                            df_tab[idx_col] = ce_err
                        else:
                            if mat != "TLD":
                                vals1 = self.raw_data[t][4]["Value"].values[: len(x)]
                            else:
                                if self.testname == "FNG-SiC":
                                    # Neutron dose
                                    Dn = (
                                        self.raw_data[t][16]["Value"].values[: len(x)]
                                    ) * fngsic_norm
                                    Dn_multiplied = [
                                        value * constant
                                        for value, constant in zip(Dn, fngsic_k)
                                    ]
                                    # Photon dose
                                    Dp = (
                                        self.raw_data[t][26]["Value"].values[: len(x)]
                                    ) * fngsic_norm
                                    # Sum neutron and photon dose with neutron sensitivity as a function of depth
                                    Dt = [sum(pair) for pair in zip(Dn_multiplied, Dp)]
                                    vals1 = Dt
                                else:
                                    vals1 = self.raw_data[t][6]["Value"].values[
                                        : len(x)
                                    ]
                            vals2 = exp_data_df.loc[:, "Reaction Rate"].to_numpy()
                            ratio = vals1 / vals2
                            ratio = ratio.tolist()
                            df_tab[idx_col] = vals1 / vals2

                # Assign worksheet title and put into Excel
                conv_df = self._get_conv_df(mat, len(x))
                sheet = self.testname.replace("-", " ")
                sheet_name = sheet + ", Foil {}".format(mat)
                df_tab.to_excel(writer, sheet_name=sheet_name)
                conv_df.to_excel(writer, sheet_name=sheet_name, startrow=18)

    def _build_atlas(self, tmp_path: str | os.PathLike, atlas: at.Atlas) -> at.Atlas:
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str | os.PathLike
            path to the temporary folder containing the plots for the atlas
        atlas : at.Atlas
            Object representing the plot Atlas.

        Returns
        -------
        atlas : at.Atlas
            Object representing the plot Atlas.
        """
        # FNG SiC specific corrections/normalisations
        fngsic_k = [0.212, 0.204, 0.202, 0.202]  # Neutron sensitivity of TL detectors
        fngsic_norm = 1.602e-13 * 1000  # J/MeV * g/kg
        # Set plot and axes details
        unit = "-"
        xlabel = "Shielding thickness [cm]"
        data = []
        # TODO Replace when other transport codes implemented.
        code = "mcnp"
        for material in tqdm(self.inputs, desc="Foil: "):
            data = []
            exp_folder = os.path.join(self.path_exp_res, material)
            exp_filename = self.testname + "_" + material + ".csv"
            exp_filepath = os.path.join(exp_folder, exp_filename)
            exp_data_df = pd.read_csv(exp_filepath)
            # Get experimental data and errors for the selected benchmark case
            x = exp_data_df["Depth"].values
            y = []
            err = []
            y.append(exp_data_df["Reaction Rate"].values)
            err.append(exp_data_df["Error"].values / 100)
            # Append experimental data to data list (sent to plotter)
            ylabel = "Experiment"
            data_exp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
            data.append(data_exp)

            if material != "TLD":
                title = self.testname + " experiment, Foil: " + FOILS_REACTION[material]
            else:
                if self.testname == "FNG-SiC":
                    title = (
                        self.testname
                        + " experiment, Total absorbed dose in TLD detectors"
                    )
                else:
                    title = (
                        self.testname
                        + " experiment, Gamma absorbed dose in TLD-300 detectors"
                    )
            # Loop over selected libraries
            # Loop over selected libraries
            for lib in self.lib[1:]:
                # Get library name, assign title to the plot
                ylabel = self.session.conf.get_lib_name(lib)
                y = []
                err = []
                if material != "TLD":
                    v = self.raw_data[(material, lib)][4]["Value"].values[: len(x)]
                else:
                    if self.testname == "FNG-SiC":
                        # Neutron dose
                        Dn = (
                            self.raw_data[(material, lib)][16]["Value"].values[: len(x)]
                        ) * fngsic_norm
                        Dn_multiplied = [
                            value * constant for value, constant in zip(Dn, fngsic_k)
                        ]
                        # Photon dose
                        Dp = (
                            self.raw_data[(material, lib)][26]["Value"].values[: len(x)]
                        ) * fngsic_norm
                        # Sum neutron and photon dose with neutron sensitivity as a function of depth
                        v = [sum(pair) for pair in zip(Dn_multiplied, Dp)]
                    else:
                        v = self.raw_data[(material, lib)][6]["Value"].values[: len(x)]
                y.append(v)
                if material != "TLD":
                    v = self.raw_data[(material, lib)][4]["Error"].values[: len(x)]
                else:
                    if self.testname == "FNG-SiC":
                        v = np.sqrt(
                            np.square(
                                self.raw_data[(material, lib)][16]["Error"].values[
                                    : len(x)
                                ]
                            )
                            + np.square(
                                self.raw_data[(material, lib)][26]["Error"].values[
                                    : len(x)
                                ]
                            )
                        )
                    else:
                        v = self.raw_data[(material, lib)][6]["Error"].values[: len(x)]
                err.append(v)
                # Append computational data to data list(to be sent to plotter)
                data_comp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                data.append(data_comp)

            # Send data to plotter
            outname = "tmp"
            if material != "TLD":
                quantity = [ACTIVATION_REACTION[material] + " Reaction Rate"]
            else:
                quantity = ["Absorbed dose"]
            atlas.doc.add_heading(title, level=1)
            plot = Plotter(
                data, title, tmp_path, outname, quantity, unit, xlabel, self.testname
            )
            img_path = plot.plot("Waves")
            atlas.insert_img(img_path, width=Inches(9))
            atlas.doc.add_page_break()

        return atlas

    def _get_conv_df(self, mat: str, size: int) -> pd.DataFrame:
        """
        Method to calculate average and maximum uncertainties

        Parameters
        ----------
        mat : str
            String denoting material
        size : int
            Integer denoting size of array

        Returns
        -------
        conv_df : pd.DataFrame
            Dataframe containing Max Error and Average Error columns
        """
        conv_df = pd.DataFrame()
        for lib in self.lib[1:]:
            if mat != "TLD":
                max = self.raw_data[(mat, lib)][4]["Error"].values[:size].max()
                avg = self.raw_data[(mat, lib)][4]["Error"].values[:size].mean()
            else:
                if self.testname == "FNG-SiC":
                    v = np.sqrt(
                        np.square(self.raw_data[(mat, lib)][16]["Error"].values[:size])
                        + np.square(
                            self.raw_data[(mat, lib)][26]["Error"].values[:size]
                        )
                    )
                    max = np.max(v)
                    avg = np.mean(v)
                else:
                    max = self.raw_data[(mat, lib)][6]["Error"].values[:size].max()
                    avg = self.raw_data[(mat, lib)][6]["Error"].values[:size].mean()
            library = self.session.conf.get_lib_name(lib)
            conv_df.loc["Max Error", library] = max
            conv_df.loc["Average Error", library] = avg
        return conv_df


class MultipleSpectrumOutput(SpectrumOutput):
    def _build_atlas(self, tmp_path: str | os.PathLike, atlas: at.Atlas) -> at.Atlas:
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str | os.PathLike
            path to the temporary folder containing the plots for the atlas
        atlas : at.Atlas
            Object representing the plot Atlas.

        Returns
        -------
        atlas : at.Atlas
            Object representing the plot Atlas.
        """
        self.tables = []
        self.groups = pd.read_excel(self.cnf_path)
        self.groups = self.groups.set_index(["Group", "Tally", "Input"])
        self.group_list = self.groups.index.get_level_values("Group").unique().tolist()
        for group in self.group_list:
            self._plot_tally_group(group, tmp_path, atlas)

        # Dump C/E table
        self._dump_ce_table()

        return atlas

    def _plot_tally_group(
        self, group: list, tmp_path: str | os.PathLike, atlas: at.Atlas
    ) -> at.Atlas:
        """
        Plots tallies for a given group of outputs and add to Atlas object

        Parameters
        ----------
        group : list
            list of groups in the experimental benchmark object, outputs are
            grouped by material, several tallies for each material/group
        tmp_path : str or os.PathLike
            path to temporary atlas plot folder
        atlas : JADE Atlas
            Atlas object

        Returns
        -------
        atlas : JADE Atlas
            adjusted Atlas object
        """
        # Extract 'Tally' and 'Input' values for the current 'Group'
        group_data = self.groups.xs(group, level="Group", drop_level=False)
        data_group = {}
        group_lab = []
        mult_factors = group_data["Multiplying factor"].values.tolist()
        for m, idx in enumerate(group_data.index.tolist()):
            tallynum = idx[1]
            input = idx[2]
            if str(tallynum) not in self.results[input, self.lib[1]].keys():
                continue
            quantity = group_data.loc[(group, tallynum, input), "Quantity"]
            particle = group_data.loc[(group, tallynum, input), "Particle"]
            add_info = group_data.loc[(group, tallynum, input), "Y Label"]
            quant_string = particle + " " + quantity + " " + add_info
            e_int = group_data.loc[(group, tallynum, input), "C/E X Quantity intervals"]
            e_int = e_int.split("-")

            # Convert the list of number strings into a list of integers
            e_intervals = [float(num) for num in e_int]
            data_temp, xlabel = self._data_collect(
                input, str(tallynum), quant_string, e_intervals
            )
            if data_temp is None:
                continue
            data_group[m] = data_temp
            unit = group_data.loc[(group, tallynum, input), "Y Unit"]

            group_lab.append(add_info)
            # Once the data is collected it is passed to the plotter
        title = self._define_title(input, particle, quantity)
        outname = "tmp"
        plot = Plotter(
            data_group,
            title,
            tmp_path,
            outname,
            quantity,
            unit,
            xlabel,
            self.testname,
            group_num=group,
            add_labels=group_lab,
            mult_factors=mult_factors,
        )
        img_path = plot.plot("Experimental points group")
        atlas.doc.add_heading(title, level=1)
        atlas.insert_img(img_path)
        img_path = plot.plot("Experimental points group CE")
        atlas.doc.add_heading(title + " C/E", level=1)
        atlas.insert_img(img_path, width=Inches(9))
        return atlas

    def _define_title(self, input: str, particle: str, quantity: str) -> str:
        """
        Determines which benchmark is being compared and assigns title
        accordinly

        Parameters
        ----------
        input : str
            Test name
        particle : str
            Particle being tallied
        quantity : str
            Type of quantity being plotted on the X axis

        Returns
        -------
        Title: str
            Title string
        """

        if not self.multiplerun:
            title = self.testname + ", " + particle + " " + quantity
        elif self.testname == "Tiara-BC":
            mat = input.split("-")[0]
            if mat == "cc":
                material = "Concrete"
            else:
                material = "Iron"
            energy = input.split("-")[1]
            sh_th = input.split("-")[2]
            add_coll = input.split("-")[3]
            title = (
                self.testname
                + ", Shielding: "
                + material
                + ", "
                + sh_th
                + "cm; Source energy: "
                + energy
                + " MeV; Additional collimator: "
                + add_coll
                + " cm"
            )
        elif self.testname == "FNS-TOF":
            mat = input.split("-")[0]
            sl_th = input.split("-")[1]
            title = self.testname + ", " + sl_th + "cm " + mat + " slab"
        else:
            title = self.testname + ", " + particle + " " + quantity
        return title


class FNGHCPBOutput(ExperimentalOutput):
    def _processMCNPdata(self, output: MCNPSimOutput) -> None:
        """
        Used to override parent function as this is not required.

        Parameters
        ----------
        output : MCNPSimOutput
            MCNP simulation output object

        Returns
        -------
        None
        """
        return None

    def _pp_excel_comparison(self) -> None:
        """
        This method prints C/E tables for shielding benchmark comparisons

        Returns
        -------
        None.
        """

        lib_names_dict = {}
        column_names = []
        column_names.append(("Exp", "Value"))
        column_names.append(("Exp", "Error"))
        for lib in self.lib[1:]:
            namelib = self.session.conf.get_lib_name(lib)
            lib_names_dict[namelib] = lib
            column_names.append((namelib, "Value"))
            column_names.append((namelib, "C/E"))
            column_names.append((namelib, "C/E Error"))

        names = ["Library", ""]
        column_index = pd.MultiIndex.from_tuples(column_names, names=names)
        filepath = self.excel_path + "\\" + self.testname + "_CE_tables.xlsx"
        with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
            code = "mcnp"
            for mat in self.inputs:
                exp_folder = os.path.join(self.path_exp_res, mat)
                exp_filename = self.testname + "_" + mat + ".csv"
                exp_filepath = os.path.join(exp_folder, exp_filename)
                exp_data_df = pd.read_csv(exp_filepath)

                # Get experimental data and errors for the selected benchmark case
                if mat == "H3":
                    x = exp_data_df["Pellet"].values.tolist()
                    indexes = pd.Index(data=x, name="Pellet #")
                else:
                    x = exp_data_df["Depth"].values.tolist()
                    indexes = pd.Index(data=x, name="Depth [cm]")

                df_tab = pd.DataFrame(index=indexes, columns=column_index)
                for idx_col in df_tab.columns.values.tolist():
                    if idx_col[0] == "Exp":
                        if idx_col[1] == "Value":
                            if mat == "H3":
                                vals = exp_data_df.loc[:, "Activity"].tolist()
                            else:
                                vals = exp_data_df.loc[:, "Reaction Rate"].tolist()
                            df_tab[idx_col] = vals
                        else:
                            vals = exp_data_df.loc[:, "Error"].to_numpy() / 100
                            vals = vals.tolist()
                            df_tab[idx_col] = vals
                    else:
                        t = (mat, lib_names_dict[idx_col[0]])
                        if idx_col[1] == "Value":
                            if mat != "H3":
                                vals = self.raw_data[t][4]["Value"].values[: len(x)]
                            else:
                                # Total activity
                                vals = []
                                for i in range(4):
                                    vals.extend(
                                        (self.raw_data[t][84]["Value"].values[i::4])
                                    )

                            df_tab[idx_col] = vals

                        elif idx_col[1] == "C/E Error":
                            if mat != "H3":
                                errs = self.raw_data[t][4]["Error"].values[: len(x)]
                            else:
                                errs = []
                                for i in range(4):
                                    yerr = self.raw_data[t][84]["Error"].values[i::4]
                                    errs.extend(yerr)

                            vals1 = np.square(errs)
                            vals2 = np.square(
                                exp_data_df.loc[:, "Error"].to_numpy() / 100
                            )
                            ce_err = np.sqrt(vals1 + vals2)
                            ce_err = ce_err.tolist()
                            df_tab[idx_col] = ce_err
                        # Calculate C/E value
                        else:
                            if mat != "H3":
                                vals1 = self.raw_data[t][4]["Value"].values[: len(x)]
                            else:
                                vals1 = []
                                for i in range(4):
                                    vals1.extend(
                                        self.raw_data[t][84]["Value"].values[i::4]
                                    )

                            if mat == "H3":
                                vals2 = exp_data_df.loc[:, "Activity"].to_numpy()
                            else:
                                vals2 = exp_data_df.loc[:, "Reaction Rate"].to_numpy()
                            ratio = vals1 / vals2
                            ratio = ratio.tolist()
                            df_tab[idx_col] = vals1 / vals2

                # Assign worksheet title and put into Excel
                conv_df = self._get_conv_df(mat, len(x))
                sheet = self.testname.replace("-", " ")
                if mat != "H3":
                    sheet_name = sheet + ", Foil {}".format(mat)
                else:
                    sheet_name = sheet + " H3 activity"
                df_tab.to_excel(writer, sheet_name=sheet_name)
                conv_df.to_excel(writer, sheet_name=sheet_name, startrow=55)
                # Close the Pandas Excel writer object and output the Excel file

    def _build_atlas(self, tmp_path: str | os.PathLike, atlas: at.Atlas) -> at.Atlas:
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str | os.PathLike
            path to the temporary folder containing the plots for the atlas
        atlas : at.Atlas
            Object representing the plot Atlas.

        Returns
        -------
        atlas : at.Atlas
            Object representing the plot Atlas.
        """
        for material in tqdm(self.inputs):
            # Tritium Activity
            if material == "H3":
                unit = "Bq/g"
                quantity = "Activity"
                for i in range(4):
                    data = []
                    # y = []
                    # err = []
                    exp_folder = os.path.join(self.path_exp_res, material)
                    exp_filename = self.testname + "_" + material + ".csv"
                    exp_filepath = os.path.join(exp_folder, exp_filename)
                    exp_data_df = pd.read_csv(exp_filepath)

                    xlabel = "Pellet no."
                    x = list(range(1, 13))

                    y = exp_data_df["Activity"].values[i * 12 : (i + 1) * 12]
                    err = exp_data_df["Error"].values[i * 12 : (i + 1) * 12] / 100

                    ylabel_exp = "Experiment"
                    data_exp = {"x": x, "y": y, "err": err, "ylabel": ylabel_exp}
                    data.append(data_exp)

                    for lib in self.lib[1:]:
                        # y = []
                        # err = []
                        # Total tritium production Li6 + Li7
                        ycalc = self.raw_data[(material, lib)][84]["Value"].values[i::4]

                        yerr = np.square(
                            self.raw_data[(material, lib)][84]["Error"].values[i::4]
                        )

                        y = ycalc
                        err = yerr

                        ylabel_calc = self.session.conf.get_lib_name(lib)
                        data_calc = {"x": x, "y": y, "err": err, "ylabel": ylabel_calc}
                        data.append(data_calc)

                    title = f"ENEA{2*(i+1)} pellet stack"
                    outname = "tmp"
                    plot = Plotter(
                        data,
                        title,
                        tmp_path,
                        outname,
                        quantity,
                        unit,
                        xlabel,
                        self.testname,
                    )
                    img_path = plot.plot("Discrete Experimental points")
                    atlas.insert_img(img_path)
            # Foils
            else:
                unit = "-"
                quantity = ["C/E"]
                data = []
                exp_folder = os.path.join(self.path_exp_res, material)
                exp_filename = self.testname + "_" + material + ".csv"
                exp_filepath = os.path.join(exp_folder, exp_filename)
                exp_data_df = pd.read_csv(exp_filepath)

                # Get experimental data and errors for the selected benchmark case
                xlabel = "Shielding thickness [cm]"
                x = list(exp_data_df["Depth"].values)
                y = []
                err = []
                y.append(exp_data_df["Reaction Rate"].values)
                err.append(exp_data_df["Error"].values / 100)
                # Append experimental data to data list (sent to plotter)
                ylabel = "Experiment"
                data_exp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                data.append(data_exp)

                title = self.testname + " experiment, Foil: " + material

                # Loop over selected libraries
                for lib in self.lib[1:]:
                    # Get library name, assign title to the plot
                    ylabel = self.session.conf.get_lib_name(lib)
                    y = []
                    err = []

                    ycalc = self.raw_data[(material, lib)][4]["Value"].values[: len(x)]
                    y.append(ycalc)

                    yerr = self.raw_data[(material, lib)][4]["Error"].values[: len(x)]
                    err.append(yerr)

                    # Append computational data to data list(to be sent to plotter)
                    data_comp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                    data.append(data_comp)

                outname = "tmp"
                plot = Plotter(
                    data,
                    title,
                    tmp_path,
                    outname,
                    quantity,
                    unit,
                    xlabel,
                    self.testname,
                )
                img_path = plot.plot("Waves")
                atlas.insert_img(img_path)
        return atlas

    def _get_conv_df(self, mat: str, size: int) -> pd.DataFrame:
        """
        Method to calculate average and maximum uncertainties

        Parameters
        ----------
        mat : str
            String denoting material
        size : int
            Integer denoting size of array

        Returns
        -------
        conv_df : pd.DataFrame
            Dataframe containing Max Error and Average Error columns
        """
        conv_df = pd.DataFrame()
        for lib in self.lib[1:]:
            if mat != "H3":
                max = self.raw_data[(mat, lib)][4]["Error"].values[:size].max()
                avg = self.raw_data[(mat, lib)][4]["Error"].values[:size].mean()
            else:
                max = self.raw_data[(mat, lib)][84]["Error"].values[:size].max()
                avg = self.raw_data[(mat, lib)][84]["Error"].values[:size].mean()
            library = self.session.conf.get_lib_name(lib)
            conv_df.loc["Max Error", library] = max
            conv_df.loc["Average Error", library] = avg
        return conv_df
